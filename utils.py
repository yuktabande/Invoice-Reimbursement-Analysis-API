import io
import os
from docx import Document
import fitz  # PyMuPDF
import logging

logger = logging.getLogger(__name__)

def extract_text_from_docx(docx_content: bytes) -> str:
    """
    Extract text from a DOCX file with enhanced structure preservation
    
    Args:
        docx_content: Bytes content of the DOCX file
        
    Returns:
        Extracted text as string with preserved formatting
    """
    try:
        # Create a file-like object from bytes
        docx_stream = io.BytesIO(docx_content)
        
        # Load the document
        doc = Document(docx_stream)
        
        # Extract text from all paragraphs with structure
        text_content = []
        
        # Process paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                # Check if paragraph looks like a heading
                if paragraph.style.name.startswith('Heading') or len(paragraph.text) < 100:
                    text_content.append(f"\n=== {paragraph.text.strip()} ===")
                else:
                    text_content.append(paragraph.text.strip())
        
        # Extract text from tables with better formatting
        for table_num, table in enumerate(doc.tables):
            text_content.append(f"\n--- TABLE {table_num + 1} ---")
            for row_num, row in enumerate(table.rows):
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                
                if row_text:
                    if row_num == 0:  # Header row
                        text_content.append("HEADERS: " + " | ".join(row_text))
                    else:
                        text_content.append(" | ".join(row_text))
        
        extracted_text = "\n".join(text_content)
        
        # Clean the text
        cleaned_text = clean_text(extracted_text)
        
        logger.info(f"Extracted {len(cleaned_text)} characters from DOCX with {len(doc.paragraphs)} paragraphs and {len(doc.tables)} tables")
        
        return cleaned_text
        
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {str(e)}")
        return ""

def extract_text_from_pdf(pdf_content: bytes) -> str:
    """
    Extract text from a PDF file using PyMuPDF with enhanced formatting
    
    Args:
        pdf_content: Bytes content of the PDF file
        
    Returns:
        Extracted text as string with preserved structure
    """
    try:
        # Create a file-like object from bytes
        pdf_stream = io.BytesIO(pdf_content)
        
        # Open PDF document
        pdf_document = fitz.open(stream=pdf_stream, filetype="pdf")
        
        text_content = []
        
        # Extract text from each page with better formatting
        for page_num in range(pdf_document.page_count):
            page = pdf_document.load_page(page_num)
            
            # Get text with layout preservation
            text = page.get_text("text")
            
            # Also try to extract table data if available
            try:
                tables = page.find_tables()
                if tables:
                    for table in tables:
                        table_data = table.extract()
                        for row in table_data:
                            if row and any(cell for cell in row if cell and str(cell).strip()):
                                text += "\n" + " | ".join([str(cell).strip() if cell else "" for cell in row])
            except:
                pass  # Continue if table extraction fails
            
            if text.strip():
                text_content.append(f"=== PAGE {page_num + 1} ===")
                text_content.append(text.strip())
        
        pdf_document.close()
        
        extracted_text = "\n\n".join(text_content)
        
        # Clean and structure the text
        cleaned_text = clean_text(extracted_text)
        
        logger.info(f"Extracted {len(cleaned_text)} characters from PDF ({pdf_document.page_count} pages)")
        
        return cleaned_text
        
    except Exception as e:
        logger.error(f"Error extracting text from PDF: {str(e)}")
        return ""

def load_prompt_template() -> str:
    """
    Load the prompt template from prompt.txt file
    
    Returns:
        Prompt template as string
    """
    try:
        prompt_file_path = os.path.join(os.path.dirname(__file__), 'prompt.txt')
        
        with open(prompt_file_path, 'r', encoding='utf-8') as file:
            template = file.read()
            
        logger.info("Prompt template loaded successfully")
        return template
        
    except FileNotFoundError:
        logger.error("prompt.txt file not found")
        # Return a fallback template
        return """You are an expert HR and finance analyst responsible for verifying reimbursement invoices based on a company's official policy document.

Input:
1. A company reimbursement policy (detailed below).
2. An employee invoice (detailed below).

Your task:
- Carefully compare the invoice contents (date, items, amount, purpose, tax, category) against the company's reimbursement policy.
- Determine whether the invoice should be Fully Reimbursed, Partially Reimbursed, or Declined.
- Explain the decision with specific references to policy rules and amounts.

Guidelines:
- If all items in the invoice are within policy rules and limits, mark as "Fully Reimbursed".
- If some items or amounts exceed policy limits but are otherwise valid, mark as "Partially Reimbursed" and give reimbursable amount.
- If the invoice contains non-reimbursable or restricted items, mark as "Declined" and give the reason.
- Reimbursable amount should always be an integer.
- Only use the rules from the provided policy. Do not assume anything not mentioned.

Format your response as JSON:
{{
  "invoice": "<invoice_filename>",
  "status": "Fully Reimbursed | Partially Reimbursed | Declined",
  "amount": <reimbursable_integer>,
  "reason": "<short explanation of decision>"
}}

--- POLICY DOCUMENT ---
{policy_text}

--- INVOICE DOCUMENT ---
{invoice_text}"""
        
    except Exception as e:
        logger.error(f"Error loading prompt template: {str(e)}")
        raise

def validate_file_content(content: bytes, file_type: str) -> bool:
    """
    Validate if the file content is valid for the specified type
    
    Args:
        content: File content as bytes
        file_type: Expected file type ('docx' or 'pdf')
        
    Returns:
        True if valid, False otherwise
    """
    try:
        if file_type.lower() == 'docx':
            # Try to open as DOCX
            docx_stream = io.BytesIO(content)
            doc = Document(docx_stream)
            return True
            
        elif file_type.lower() == 'pdf':
            # Try to open as PDF
            pdf_stream = io.BytesIO(content)
            pdf_document = fitz.open(stream=pdf_stream, filetype="pdf")
            is_valid = pdf_document.page_count > 0
            pdf_document.close()
            return is_valid
            
        return False
        
    except Exception as e:
        logger.error(f"File validation failed for {file_type}: {str(e)}")
        return False

def clean_text(text: str) -> str:
    """
    Clean and normalize extracted text
    
    Args:
        text: Raw extracted text
        
    Returns:
        Cleaned text
    """
    if not text:
        return ""
    
    # Remove excessive whitespace
    lines = [line.strip() for line in text.split('\n') if line.strip()]
    
    # Join lines with single newlines
    cleaned_text = '\n'.join(lines)
    
    # Remove multiple consecutive newlines
    while '\n\n\n' in cleaned_text:
        cleaned_text = cleaned_text.replace('\n\n\n', '\n\n')
    
    return cleaned_text

def extract_key_info_from_invoice(invoice_text: str) -> dict:
    """
    Extract key information from invoice text (optional utility)
    
    Args:
        invoice_text: Extracted invoice text
        
    Returns:
        Dictionary with extracted key information
    """
    key_info = {
        "has_date": False,
        "has_amount": False,
        "has_vendor": False,
        "estimated_amount": None
    }
    
    lines = invoice_text.lower().split('\n')
    
    # Look for common patterns
    for line in lines:
        # Check for date patterns
        if any(word in line for word in ['date', 'dated', 'invoice date']):
            key_info["has_date"] = True
        
        # Check for amount patterns
        if any(symbol in line for symbol in ['$', '₹', 'rs', 'amount', 'total']):
            key_info["has_amount"] = True
            
        # Check for vendor information
        if any(word in line for word in ['vendor', 'company', 'ltd', 'inc', 'pvt']):
            key_info["has_vendor"] = True
    
    return key_info